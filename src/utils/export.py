from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pdfkit
from typing import Dict, Any, Optional, List  # Added List import
from datetime import datetime
import os
from io import BytesIO
import pdfkit
import tempfile
from pathlib import Path
import streamlit as st
import subprocess

def convert_to_pdf(content: str, output_path: str) -> str:
    """
    Convert content to PDF using pandoc.
    
    Args:
        content (str): The content to convert
        output_path (str): Path where the PDF file should be saved
    
    Returns:
        str: Path to the created PDF file
    """
    # Create temporary HTML file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as temp_html:
        html_content = f"""
        <!DOCTYPE html>
        <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        font-size: 12pt;
                        line-height: 1.5;
                        margin: 1in;
                    }}
                </style>
            </head>
            <body>
                {content.replace(chr(10), '<br>')}
            </body>
        </html>
        """
        temp_html.write(html_content)
        temp_html_path = temp_html.name

    try:
        # Convert HTML to PDF using pandoc
        subprocess.run([
            'pandoc',
            temp_html_path,
            '-o', output_path,
            '--pdf-engine=wkhtmltopdf',
            '-V', 'geometry:margin=1in'
        ], check=True)
        
        return output_path
    except subprocess.CalledProcessError as e:
        raise Exception(f"Error converting to PDF: {str(e)}")
    finally:
        # Clean up temporary file
        Path(temp_html_path).unlink(missing_ok=True)

class ExportManager:
    """Manages document export operations."""
    
    def __init__(self):
        self.doc = None
    
    def export_to_docx(self, content: str, doc_type: str) -> Optional[bytes]:
        """
        Export content to DOCX format.
        
        Args:
            content (str): Content to export
            doc_type (str): Type of document ('cover_letter' or 'resume')
            
        Returns:
            bytes: DOCX file as bytes or None if failed
        """
        try:
            doc = Document()
            
            # Configure document
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add content
            doc.add_paragraph(content)
            
            # Save to bytes
            docx_bytes = BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            return docx_bytes.getvalue()
            
        except Exception as e:
            print(f"Error exporting to DOCX: {str(e)}")
            return None
    
    def export_to_pdf(self, content: str) -> Optional[bytes]:
        """
        Export content to PDF format.
        
        Args:
            content (str): Content to export
            
        Returns:
            bytes: PDF file as bytes or None if failed
        """
        try:
            # Create temporary HTML file
            with open('temp.html', 'w', encoding='utf-8') as f:
                f.write(f"<html><body>{content}</body></html>")
            
            # Convert to PDF
            pdf_bytes = BytesIO()
            pdfkit.from_file('temp.html', pdf_bytes)
            
            # Clean up
            os.remove('temp.html')
            
            return pdf_bytes.getvalue()
            
        except Exception as e:
            print(f"Error exporting to PDF: {str(e)}")
            return None

def create_cover_letter_docx(content: str, output_path: str) -> str:
    """Create a DOCX file with proper formatting and contact information."""
    doc = Document()
    
    # Set the default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Add contact information header if available
    if hasattr(st.session_state, 'additional_info'):
        info = st.session_state.additional_info
        
        # Contact Header
        if info.get('full_name'):
            p = doc.add_paragraph(info['full_name'])
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
        if info.get('email') or info.get('phone'):
            contact_line = []
            if info.get('email'):
                contact_line.append(info['email'])
            if info.get('phone'):
                contact_line.append(info['phone'])
            if contact_line:
                doc.add_paragraph(' | '.join(contact_line))
                
        if info.get('location'):
            doc.add_paragraph(info['location'])
            
        if info.get('linkedin'):
            doc.add_paragraph(info['linkedin'])
            
        # Add date
        doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        
        # Add spacing
        doc.add_paragraph()
        
        # Company Header if available
        if info.get('company_name') or info.get('hiring_manager'):
            if info.get('hiring_manager'):
                doc.add_paragraph(f"Dear {info['hiring_manager']},")
            else:
                doc.add_paragraph("Dear Hiring Manager,")
    
    # Add main content
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph(para.strip())
            # Ensure each paragraph uses the correct font settings
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
    
    # Save the document
    doc.save(output_path)
    return output_path

def convert_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """
    Convert DOCX to PDF using pdfkit.
    
    Args:
        docx_path (str): Path to source DOCX file
        pdf_path (str): Path to save PDF file
    
    Returns:
        bool: Success status
    """
    try:
        # Convert DOCX to HTML first (simplified approach)
        os.system(f'pandoc -f docx -t html "{docx_path}" -o temp.html')
        
        # Convert HTML to PDF
        pdfkit.from_file('temp.html', pdf_path)
        
        # Clean up temporary file
        os.remove('temp.html')
        return True
        
    except Exception as e:
        print(f"Error converting to PDF: {str(e)}")
        return False

def export_resume_docx(content: Dict[str, Any], output_path: str) -> bool:
    """
    Create a formatted resume DOCX file.
    
    Args:
        content (Dict[str, Any]): Resume content and sections
        output_path (str): Path to save the document
    
    Returns:
        bool: Success status
    """
    try:
        doc = Document()
        
        # Add name and contact info
        name = doc.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.add_run(content.get('name', '')).bold = True
        
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(f"{content.get('email', '')} | {content.get('phone', '')}")
        
        # Add sections
        for section in content.get('sections', []):
            # Add section heading
            doc.add_paragraph()
            heading = doc.add_paragraph()
            heading.add_run(section['title'].upper()).bold = True
            
            # Add section content
            if isinstance(section['content'], list):
                for item in section['content']:
                    p = doc.add_paragraph(item, style='List Bullet')
            else:
                doc.add_paragraph(section['content'])
        
        # Save document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating resume DOCX: {str(e)}")
        return False

def generate_exports(content: Dict[str, Any], base_filename: str, formats: List[str]) -> Dict[str, str]:
    """
    Generate exports in specified formats.
    
    Args:
        content: Dictionary containing document content
        base_filename: Base name for output files
        formats: List of formats to generate ("DOCX", "PDF", or "Both")
    
    Returns:
        Dictionary with paths to generated files
    """
    results = {}
    
    try:
        # Create temp directory if it doesn't exist
        temp_dir = Path("temp")
        temp_dir.mkdir(exist_ok=True)
        
        # Generate DOCX
        if "DOCX" in formats or "Both" in formats:
            docx_path = temp_dir / f"{base_filename}.docx"
            create_cover_letter_docx(content, str(docx_path))
            results["docx"] = str(docx_path)
        
        # Generate PDF
        if "PDF" in formats or "Both" in formats:
            pdf_path = temp_dir / f"{base_filename}.pdf"
            if "docx" in results:
                convert_to_pdf(results["docx"], str(pdf_path))
            results["pdf"] = str(pdf_path)
        
        return results
    
    except Exception as e:
        print(f"Error generating exports: {str(e)}")
        return {}